from datetime import datetime, timedelta
import os
from docx import Document
from watcalendars.utils.logutils import log_entry, log, ERROR, OK
from watcalendars.utils.config import BLOCK_TIMES, ROMAN_MONTH, DATE_TOKEN_RE, TYPE_FULL_MAP
import re

TIME_RE = re.compile(r"(\d{1,2}):(\d{2})")

DAY_ALIASES = {
    'pon.': 0, 'wt.': 1, 'śr.': 2, 'sr.': 2, 'czw.': 3, 'pt.': 4, 'sob.': 5, 'niedz.': 6
}

def _pick_base_year(doc: Document, filepath: str) -> int:
    """Determine the academic autumn year for the schedule.
    Priority:
    1) Parse from filename like '.../WIG24KX1S1.docx' -> 2024
    2) Parse from any 'dd.mm.yyyy' in document paragraphs (uses that year as autumn year)
    3) Fallback to current year (treated as autumn year)
    """
    try:
        fname = os.path.basename(filepath)
        m = re.search(r"(?i)WIG(\d{2})", fname)
        if m:
            yy = int(m.group(1))
            return 2000 + yy
    except Exception:
        pass

    for p in doc.paragraphs:
        txt = p.text
        m = re.search(r"(\d{2})[.](\d{2})[.](\d{4})", txt)
        if m:
            return int(m.group(3))

    return datetime.now().year

def _combine_date(autumn_year: int, day: int, roman_month: str) -> datetime:
    """Map IX–XII to autumn_year, and I–VI (and beyond) to autumn_year+1."""
    month = ROMAN_MONTH.get(roman_month, 1)
    year = autumn_year if month >= 9 else autumn_year + 1
    return datetime(year, month, day)

def _parse_cell_entry(text: str):
    text = text.strip()
    subject = text
    type_symbol = ""
    type_full = ""
    room = ""
    lecturer = ""
    mtype = re.search(r"\(([^)]+)\)", text)
    if mtype:
        type_symbol = f"({mtype.group(1)})"
        type_full = TYPE_FULL_MAP.get(type_symbol, type_symbol)
    parts = [p.strip() for p in text.split('/') if p.strip()]
    if parts:
        subject = re.sub(r"\s*\([^)]*\)", "", parts[0]).strip()
    if len(parts) >= 2:
        room = parts[1]
    if len(parts) >= 3:
        lecturer = parts[2]
    return subject, type_symbol, type_full, room, lecturer

def parse_wig_docx(filepath):
    """
    Parse a WIG Word schedule file organized as a grid with per-day date headers and block times.
    Enrich events with full subject names and full lecturer names (legend columns).
    """

    def parse_log():
        logs = []
        try:
            doc = Document(filepath)
        except Exception as e:
            log_entry(f"{ERROR} Failed to open DOCX: {e}", logs)
            return []

        base_year = _pick_base_year(doc, filepath)
        lessons = []

        main_table = None
        for t in doc.tables:
            if len(t.rows) >= 10 and len(t.columns) >= 10:
                main_table = t
                break
        if main_table is None:
            if doc.tables:
                main_table = doc.tables[0]
            else:
                log_entry(f"{ERROR} No tables in DOCX", logs)
                return []

        subject_legend_full: dict[str, str] = {}
        lecturer_names_pool: list[str] = []

        def is_lecturer_name(text: str) -> bool:
            return re.search(r"(?i)\b(prof\.|dr|hab\.|mgr|inż\.|inz\.|ppłk|płk|mjr|kpt\.|por\.|ppor\.|chor\.|sierż\.|kpr\.)\b", text or "") is not None

        def collect_legends(table):
            cols = len(table.columns)
            if cols < 2:
                return
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue
                left = cells[-2].strip()
                right = cells[-1].strip()
                if not right:
                    continue
                if is_lecturer_name(right):
                    if right not in lecturer_names_pool:
                        lecturer_names_pool.append(right)
                    continue
                if left and re.fullmatch(r"[A-Za-zĄĆĘŁŃÓŚŹŻ0-9]{2,8}", left) and len(right) >= 5:
                    subject_legend_full.setdefault(left, right)

        collect_legends(main_table)

        import unicodedata
        def undiac(s: str) -> str:
            return ''.join(c for c in unicodedata.normalize('NFD', s or '') if unicodedata.category(c) != 'Mn')

        def pick_lecturers_for_code(code: str) -> list[str]:
            names = lecturer_names_pool
            if not names:
                return []
            token = re.sub(r'[^a-z]', '', undiac((code or '').lower()))
            if not token:
                return []
            result: list[str] = []
            seen: set[str] = set()
            norm_names = [(full, re.sub(r'[^a-z]', '', undiac(full).lower())) for full in names]
            for full, hay in norm_names:
                if full in seen:
                    continue
                if all(ch in hay for ch in set(token)):
                    result.append(full)
                    seen.add(full)
            return result[:2]

        current_date_cols: dict[int, datetime] = {}

        for row in main_table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells or len(cells) < 2:
                continue
            first = (cells[0] or '').lower()

            if first in DAY_ALIASES and len(cells) >= 3:
                temp: dict[int, datetime] = {}
                count_dates = 0
                last_idx = -1
                for j, val in enumerate(cells[2:], start=2):
                    m = DATE_TOKEN_RE.match(val)
                    if m:
                        d = int(m.group(1))
                        roman = m.group(2)
                        temp[j] = _combine_date(base_year, d, roman)
                        count_dates += 1
                        last_idx = j
                if count_dates >= 1:
                    current_date_cols = {j: d for j, d in temp.items() if j <= last_idx}
                    continue

            block_token = cells[1]
            if block_token not in BLOCK_TIMES or not current_date_cols:
                continue
            start_time_str, end_time_str = BLOCK_TIMES[block_token]

            for cj, date_obj in current_date_cols.items():
                if cj >= len(cells):
                    continue
                cell_text = cells[cj].strip()
                if not cell_text:
                    continue
                raw = cell_text.replace('\r', '\n')
                lines = [ln.strip() for ln in raw.split('\n') if ln.strip()]
                entries = []
                if len(lines) >= 2 and any(re.search(r"\([^)]*\)", ln) for ln in lines):
                    entries = [" / ".join(lines)]
                else:
                    if ';' in raw:
                        entries = [e.strip() for e in raw.split(';') if e.strip()]
                    else:
                        entries = [raw.strip()]
                for entry in entries:
                    subject, type_symbol, type_full, room, lecturer_code = _parse_cell_entry(entry)
                    full_subject = subject_legend_full.get(subject, subject)
                    lecturers = pick_lecturers_for_code(lecturer_code)
                    sh, sm = map(int, start_time_str.split(':'))
                    eh, em = map(int, end_time_str.split(':'))
                    start_dt = date_obj.replace(hour=sh, minute=sm, second=0, microsecond=0)
                    end_dt = date_obj.replace(hour=eh, minute=em, second=0, microsecond=0)
                    lessons.append({
                        "start": start_dt,
                        "end": end_dt,
                        "subject": subject,
                        "type": type_symbol or type_full,
                        "type_full": type_full or type_symbol,
                        "room": room,
                        "lecturer": "; ".join(lecturers) if lecturers else "",
                        "lecturers": lecturers,
                        "full_subject": full_subject
                    })

        if not lessons:
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for p in paras:
                m = re.search(r"(\d{1,2}:\d{2})\s*[-–]\s*(\d{1,2}:\d{2})", p)
                if not m:
                    continue
                start_str, end_str = m.groups()
                def to_dt(tstr):
                    h, mi = map(int, tstr.split(":"))
                    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
                    return today + timedelta(hours=h, minutes=mi)
                start = to_dt(start_str)
                end = to_dt(end_str)
                subject = p
                lessons.append({
                    "start": start,
                    "end": end,
                    "subject": subject,
                    "type": "",
                    "room": "",
                    "lecturer": ""
                })

        log_entry(f"{OK} Parsed {len(lessons)} lessons from DOCX.", logs)
        return lessons

    return log(f"Parsing WIG DOCX: {filepath}", parse_log)